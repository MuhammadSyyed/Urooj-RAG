from __future__ import annotations
import json
from pathlib import Path
from typing import Dict, List

from docx import Document


ARTIFACTS_DIR = Path("../artifacts")
REPORT_MD_PATH = ARTIFACTS_DIR / "rag_llm_comparison_report.md"
REPORT_DOCX_PATH = ARTIFACTS_DIR / "rag_llm_comparison_report.docx"


def load_jsonl(path: Path) -> List[Dict]:
    """Load a file containing one JSON object per entry.

    Supports both true JSONL (one compact object per line) and
    pretty-printed objects written with json.dumps(..., indent=4).
    """
    items: List[Dict] = []
    buffer = ""

    with path.open("r", encoding="utf-8") as f:
        for raw_line in f:
            line = raw_line.strip()
            if not line:
                # skip empty lines but try to parse any accumulated buffer
                continue

            # Try simple case: a full JSON object on this line
            if not buffer:
                try:
                    items.append(json.loads(line))
                    continue
                except json.JSONDecodeError:
                    # Start accumulating a multi-line JSON object
                    buffer = raw_line
                    continue

            # We are inside a multi-line object; accumulate and try to parse
            buffer += raw_line
            try:
                items.append(json.loads(buffer))
                buffer = ""
            except json.JSONDecodeError:
                # Not complete yet, keep accumulating
                continue

    return items


def summarize_eval(eval_path: Path) -> Dict:
    eval_items = load_jsonl(eval_path)
    if not eval_items:
        return {
            "avg_score": 0.0,
            "num_items": 0,
            "hallucination_rate": 0.0,
        }

    scores = [int(it.get("score", 0)) for it in eval_items]
    avg_score = sum(scores) / len(scores)

    # Treat low scores (<=2) as hallucinations / serious errors
    hallucinations = sum(1 for s in scores if s <= 2)
    hallucination_rate = hallucinations / len(scores)

    return {
        "avg_score": avg_score,
        "num_items": len(scores),
        "hallucination_rate": hallucination_rate,
    }


def write_report(results: List[Dict]) -> None:
    # Markdown report
    lines: List[str] = []
    lines.append("# RAG + LLM Evaluation Report")
    lines.append("")
    lines.append("This report compares LLMs on the same QA dataset using an LLM-as-a-judge score (1–5).")
    lines.append("")
    lines.append("| Model | Variant | Avg Score (1-5) | Hallucination Rate | Items |")
    lines.append("|-------|---------|-----------------|---------------------|-------|")

    for r in results:
        lines.append(
            f"| {r['model_name']} | {r['variant']} | "
            f"{r['avg_score']:.2f} | {r['hallucination_rate']*100:.1f}% | {r['num_items']} |"
        )

    lines.append("")
    lines.append("## Conclusions")
    lines.append("")

    by_model: Dict[str, Dict[str, Dict]] = {}
    for r in results:
        by_model.setdefault(r["model_name"], {})[r["variant"]] = r

    for model, variants in by_model.items():
        default = variants.get("default")
        tuned = variants.get("tuned")
        if not default or not tuned:
            continue
        lines.append(f"### {model}")
        lines.append("")

        if tuned["avg_score"] > default["avg_score"]:
            lines.append("- Tuned parameters improved accuracy compared to default.")
        elif tuned["avg_score"] < default["avg_score"]:
            lines.append("- Tuned parameters reduced accuracy compared to default.")
        else:
            lines.append("- Tuned parameters kept accuracy roughly the same.")

        if tuned["hallucination_rate"] < default["hallucination_rate"]:
            lines.append("- Tuned parameters reduced hallucinations (low-score answers).")
        elif tuned["hallucination_rate"] > default["hallucination_rate"]:
            lines.append("- Tuned parameters increased hallucinations.")
        else:
            lines.append("- Tuned parameters did not significantly change hallucination rate.")

        lines.append("")

    REPORT_MD_PATH.parent.mkdir(parents=True, exist_ok=True)
    REPORT_MD_PATH.write_text("\n".join(lines), encoding="utf-8")
    print(f"Markdown report written to {REPORT_MD_PATH}")

    # DOCX report (same content, structured)
    doc = Document()
    doc.add_heading("RAG + LLM Evaluation Report", level=1)
    doc.add_paragraph(
        "This report compares multiple LLMs on the same QA dataset "
        "using an LLM-as-a-judge score (1–5)."
    )

    table = doc.add_table(rows=1, cols=5)
    hdr = table.rows[0].cells
    hdr[0].text = "Model"
    hdr[1].text = "Variant"
    hdr[2].text = "Avg Score (1-5)"
    hdr[3].text = "Hallucination Rate"
    hdr[4].text = "Items"

    for r in results:
        row = table.add_row().cells
        row[0].text = str(r["model_name"])
        row[1].text = str(r["variant"])
        row[2].text = f"{r['avg_score']:.2f}"
        row[3].text = f"{r['hallucination_rate']*100:.1f}%"
        row[4].text = str(r["num_items"])

    doc.add_page_break()
    doc.add_heading("Conclusions", level=2)

    for model, variants in by_model.items():
        default = variants.get("default")
        tuned = variants.get("tuned")
        if not default or not tuned:
            # For now, your data only has default runs; tuned can be added later.
            continue
        doc.add_heading(model, level=3)

        if tuned["avg_score"] > default["avg_score"]:
            doc.add_paragraph("Tuned parameters improved accuracy compared to default.", style="List Bullet")
        elif tuned["avg_score"] < default["avg_score"]:
            doc.add_paragraph("Tuned parameters reduced accuracy compared to default.", style="List Bullet")
        else:
            doc.add_paragraph("Tuned parameters kept accuracy roughly the same.", style="List Bullet")

        if tuned["hallucination_rate"] < default["hallucination_rate"]:
            doc.add_paragraph("Tuned parameters reduced hallucinations (low-score answers).", style="List Bullet")
        elif tuned["hallucination_rate"] > default["hallucination_rate"]:
            doc.add_paragraph("Tuned parameters increased hallucinations.", style="List Bullet")
        else:
            doc.add_paragraph("Tuned parameters did not significantly change hallucination rate.", style="List Bullet")

    REPORT_DOCX_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(REPORT_DOCX_PATH)
    print(f"DOCX report written to {REPORT_DOCX_PATH}")


def main() -> None:
    # Map of your already-run eval files: model -> {variant -> path}
    eval_files: Dict[str, Dict[str, Path]] = {
        "gemma3:1b": {
            "default": ARTIFACTS_DIR / "gemma3_eval.jsonl",
        },
        "qwen:latest": {
            "default": ARTIFACTS_DIR / "qwen_eval.jsonl",
        },
        "Phi3:latest": {
            "default": ARTIFACTS_DIR / "phi3_eval.jsonl",
        },
        "deepseek-r1:1.5b": {
            "default": ARTIFACTS_DIR / "deepseek_eval.jsonl",
        },
        "smollm:latest": {
            "default": ARTIFACTS_DIR / "smolLM_eval.jsonl",
        },
    }

    results: List[Dict] = []
    for model_name, variants in eval_files.items():
        for variant, path in variants.items():
            if not path.exists():
                continue
            summary = summarize_eval(path)
            summary.update({
                "model_name": model_name,
                "variant": variant,
            })
            results.append(summary)

    write_report(results)


if __name__ == "__main__":
    main()

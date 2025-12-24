from __future__ import annotations
import json
from pathlib import Path
from typing import Dict, List
from datetime import datetime
import io

import matplotlib.pyplot as plt
import matplotlib
matplotlib.use('Agg')  # Use non-interactive backend

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


ARTIFACTS_DIR = Path("../artifacts")
REPORT_MD_PATH = ARTIFACTS_DIR / "rag_llm_comprehensive_report.md"
REPORT_DOCX_PATH = ARTIFACTS_DIR / "rag_llm_comprehensive_report.docx"
QA_PATH = ARTIFACTS_DIR / "qa_dataset.jsonl"
GRAPHS_DIR = ARTIFACTS_DIR / "graphs"

# Ensure graphs directory exists
GRAPHS_DIR.mkdir(parents=True, exist_ok=True)


def load_jsonl(path: Path) -> List[Dict]:
    """Load a file containing one JSON object per entry.
    
    Supports both true JSONL (one compact object per line) and
    pretty-printed objects written with json.dumps(..., indent=4).
    """
    items: List[Dict] = []
    buffer = ""

    with path.open("r", encoding="utf-8") as f:
        for raw_line in f:
            line = raw_line.strip()
            if not line:
                continue

            if not buffer:
                try:
                    items.append(json.loads(line))
                    continue
                except json.JSONDecodeError:
                    buffer = raw_line
                    continue

            buffer += raw_line
            try:
                items.append(json.loads(buffer))
                buffer = ""
            except json.JSONDecodeError:
                continue

    return items


def summarize_eval(eval_path: Path, pred_path: Path = None) -> Dict:
    """Summarize an eval file, optionally including latency from predictions."""
    eval_items = load_jsonl(eval_path)
    if not eval_items:
        return {
            "avg_score": 0.0,
            "num_items": 0,
            "hallucination_rate": 0.0,
            "avg_latency": 0.0,
            "score_distribution": {},
        }

    scores = [int(it.get("score", 0)) for it in eval_items]
    avg_score = sum(scores) / len(scores) if scores else 0.0

    hallucinations = sum(1 for s in scores if s <= 2)
    hallucination_rate = hallucinations / len(scores) if scores else 0.0

    # Score distribution for reporting
    score_dist = {}
    for score in range(1, 6):
        count = sum(1 for s in scores if s == score)
        score_dist[score] = count

    result = {
        "avg_score": avg_score,
        "num_items": len(scores),
        "hallucination_rate": hallucination_rate,
        "avg_latency": 0.0,
        "score_distribution": score_dist,
    }

    # Try to get latency from predictions if available
    if pred_path and pred_path.exists():
        try:
            pred_items = load_jsonl(pred_path)
            latencies = [float(it.get("latency_sec", 0.0)) for it in pred_items if "latency_sec" in it]
            if latencies:
                result["avg_latency"] = sum(latencies) / len(latencies)
        except Exception:
            pass

    return result


def get_dataset_stats() -> Dict:
    """Get statistics about the QA dataset."""
    qa_items = load_jsonl(QA_PATH)
    if not qa_items:
        return {"total_qa_pairs": 0, "unique_docs": 0}

    docs = set()
    for item in qa_items:
        doc_id = item.get("doc_id", "unknown")
        docs.add(doc_id)

    return {
        "total_qa_pairs": len(qa_items),
        "unique_docs": len(docs),
    }


def generate_graphs(results: List[Dict]) -> Dict[str, Path]:
    """Generate comparison graphs and save as PNG files."""
    graph_paths = {}

    # 1. Average Score Comparison (Bar chart)
    fig, ax = plt.subplots(figsize=(10, 6))
    models = [r["model_name"] for r in sorted(results, key=lambda x: x["avg_score"], reverse=True)]
    scores = [r["avg_score"] for r in sorted(results, key=lambda x: x["avg_score"], reverse=True)]
    colors = ["#2ecc71" if s >= 3.5 else "#f39c12" if s >= 3.0 else "#e74c3c" for s in scores]
    
    ax.bar(models, scores, color=colors, edgecolor="black", linewidth=1.5)
    ax.set_ylabel("Average Score (1-5)", fontsize=12, fontweight="bold")
    ax.set_xlabel("Model", fontsize=12, fontweight="bold")
    ax.set_title("Model Accuracy Comparison (LLM-as-Judge Score)", fontsize=14, fontweight="bold")
    ax.set_ylim([0, 5])
    ax.axhline(y=3.0, color="gray", linestyle="--", alpha=0.5, label="Threshold (3.0)")
    ax.legend()
    plt.xticks(rotation=45, ha="right")
    plt.tight_layout()
    
    score_chart_path = GRAPHS_DIR / "01_accuracy_comparison.png"
    fig.savefig(score_chart_path, dpi=300, bbox_inches="tight")
    plt.close()
    graph_paths["accuracy"] = score_chart_path

    # 2. Hallucination Rate Comparison (Bar chart)
    fig, ax = plt.subplots(figsize=(10, 6))
    models_hal = [r["model_name"] for r in sorted(results, key=lambda x: x["hallucination_rate"])]
    hal_rates = [r["hallucination_rate"] * 100 for r in sorted(results, key=lambda x: x["hallucination_rate"])]
    colors_hal = ["#2ecc71" if h <= 15 else "#f39c12" if h <= 30 else "#e74c3c" for h in hal_rates]
    
    ax.bar(models_hal, hal_rates, color=colors_hal, edgecolor="black", linewidth=1.5)
    ax.set_ylabel("Hallucination Rate (%)", fontsize=12, fontweight="bold")
    ax.set_xlabel("Model", fontsize=12, fontweight="bold")
    ax.set_title("Hallucination Rate Comparison", fontsize=14, fontweight="bold")
    ax.set_ylim([0, 100])
    ax.axhline(y=15, color="green", linestyle="--", alpha=0.5, label="Low Risk (15%)")
    ax.axhline(y=30, color="orange", linestyle="--", alpha=0.5, label="High Risk (30%)")
    ax.legend()
    plt.xticks(rotation=45, ha="right")
    plt.tight_layout()
    
    hallucination_chart_path = GRAPHS_DIR / "02_hallucination_comparison.png"
    fig.savefig(hallucination_chart_path, dpi=300, bbox_inches="tight")
    plt.close()
    graph_paths["hallucination"] = hallucination_chart_path

    # 3. Score Distribution (Stacked bar chart)
    fig, ax = plt.subplots(figsize=(12, 6))
    models_dist = [r["model_name"] for r in sorted(results, key=lambda x: x["avg_score"], reverse=True)]
    score_1 = [r.get("score_distribution", {}).get(1, 0) for r in sorted(results, key=lambda x: x["avg_score"], reverse=True)]
    score_2 = [r.get("score_distribution", {}).get(2, 0) for r in sorted(results, key=lambda x: x["avg_score"], reverse=True)]
    score_3 = [r.get("score_distribution", {}).get(3, 0) for r in sorted(results, key=lambda x: x["avg_score"], reverse=True)]
    score_4 = [r.get("score_distribution", {}).get(4, 0) for r in sorted(results, key=lambda x: x["avg_score"], reverse=True)]
    score_5 = [r.get("score_distribution", {}).get(5, 0) for r in sorted(results, key=lambda x: x["avg_score"], reverse=True)]
    
    x = range(len(models_dist))
    ax.bar(x, score_5, label="Score 5", color="#2ecc71")
    ax.bar(x, score_4, bottom=score_5, label="Score 4", color="#3498db")
    ax.bar(x, score_3, bottom=[score_5[i] + score_4[i] for i in range(len(score_4))], label="Score 3", color="#f39c12")
    bottom_2_3 = [score_5[i] + score_4[i] + score_3[i] for i in range(len(score_3))]
    ax.bar(x, score_2, bottom=bottom_2_3, label="Score 2", color="#e67e22")
    bottom_1_2_3 = [bottom_2_3[i] + score_2[i] for i in range(len(score_2))]
    ax.bar(x, score_1, bottom=bottom_1_2_3, label="Score 1 (Hallucination)", color="#e74c3c")
    
    ax.set_ylabel("Number of Items", fontsize=12, fontweight="bold")
    ax.set_xlabel("Model", fontsize=12, fontweight="bold")
    ax.set_title("Score Distribution Across Models", fontsize=14, fontweight="bold")
    ax.set_xticks(x)
    ax.set_xticklabels(models_dist, rotation=45, ha="right")
    ax.legend(loc="upper right")
    plt.tight_layout()
    
    distribution_chart_path = GRAPHS_DIR / "03_score_distribution.png"
    fig.savefig(distribution_chart_path, dpi=300, bbox_inches="tight")
    plt.close()
    graph_paths["distribution"] = distribution_chart_path

    # 4. Accuracy vs Hallucination Scatter Plot
    fig, ax = plt.subplots(figsize=(10, 7))
    for r in results:
        ax.scatter(r["hallucination_rate"] * 100, r["avg_score"], s=300, alpha=0.7, edgecolors="black", linewidth=1.5)
        ax.annotate(r["model_name"], (r["hallucination_rate"] * 100, r["avg_score"]), 
                   fontsize=9, ha="center", va="center", fontweight="bold")
    
    ax.set_xlabel("Hallucination Rate (%)", fontsize=12, fontweight="bold")
    ax.set_ylabel("Average Accuracy Score (1-5)", fontsize=12, fontweight="bold")
    ax.set_title("Accuracy vs. Hallucination Rate Trade-off", fontsize=14, fontweight="bold")
    ax.set_xlim([-5, 105])
    ax.set_ylim([0, 5.5])
    ax.grid(True, alpha=0.3, linestyle="--")
    
    # Add quadrant lines
    ax.axhline(y=3.0, color="gray", linestyle="--", alpha=0.5)
    ax.axvline(x=15, color="gray", linestyle="--", alpha=0.5)
    ax.text(75, 5, "High Hallucination\nLower Accuracy", ha="center", fontsize=9, style="italic", alpha=0.6)
    ax.text(5, 0.5, "Low Hallucination\nLower Accuracy", ha="center", fontsize=9, style="italic", alpha=0.6)
    
    plt.tight_layout()
    scatter_chart_path = GRAPHS_DIR / "04_accuracy_vs_hallucination.png"
    fig.savefig(scatter_chart_path, dpi=300, bbox_inches="tight")
    plt.close()
    graph_paths["scatter"] = scatter_chart_path

    return graph_paths


def create_markdown_report(results: List[Dict], dataset_stats: Dict) -> str:
    """Generate comprehensive markdown report."""
    lines = []

    # Title and metadata
    lines.append("# RAG + LLM Evaluation Report: Comprehensive Comparison")
    lines.append("")
    lines.append(f"**Date:** {datetime.now().strftime('%B %d, %Y')}")
    lines.append("")

    # 1. Introduction
    lines.append("## 1. Introduction")
    lines.append("")
    lines.append(
        "This report presents a comprehensive evaluation of multiple Large Language Models (LLMs) "
        "in a Retrieval-Augmented Generation (RAG) pipeline. The goal is to compare model performance "
        "in terms of accuracy (via LLM-as-a-judge scoring), hallucination rates, and execution time. "
        "This analysis helps identify which models provide the best balance between accuracy and "
        "reliability for production RAG deployments."
    )
    lines.append("")

    # 2. Dataset Description
    lines.append("## 2. Dataset Description")
    lines.append("")
    lines.append(
        f"The evaluation uses a curated QA dataset derived from financial and banking domain documents. "
        f"Key statistics:"
    )
    lines.append("")
    lines.append(f"- **Total QA Pairs:** {dataset_stats.get('total_qa_pairs', 'N/A')}")
    lines.append(f"- **Unique Documents:** {dataset_stats.get('unique_docs', 'N/A')}")
    lines.append(f"- **Domain:** Banking, Insurance, Loans, and Financial Services")
    lines.append(f"- **Question Types:** Factual recall, comparison, definition, numerical extraction")
    lines.append("")
    lines.append(
        "Each QA pair includes a reference answer manually curated from the source documents, "
        "ensuring ground-truth labels for evaluation."
    )
    lines.append("")

    # 3. Method Description
    lines.append("## 3. Method Description")
    lines.append("")
    lines.append("### 3.1 Evaluation Methodology")
    lines.append("")
    lines.append(
        "We employ an **LLM-as-a-Judge** evaluation approach, where a secondary LLM scores each model's "
        "answer against the reference answer on a scale of 1–5:"
    )
    lines.append("")
    lines.append("- **5:** Fully correct and complete")
    lines.append("- **4:** Mostly correct, minor omissions or imprecision")
    lines.append("- **3:** Partially correct, key ideas present but incomplete")
    lines.append("- **2:** Mostly incorrect, minimal relevant content")
    lines.append("- **1:** Incorrect or irrelevant")
    lines.append("")
    lines.append(
        "Answers with scores ≤ 2 are classified as **hallucinations** or serious errors, "
        "as they deviate significantly from the reference."
    )
    lines.append("")

    lines.append("### 3.2 Key Metrics")
    lines.append("")
    lines.append("- **Average Score (1–5):** Mean LLM-as-a-judge score across all QA pairs.")
    lines.append("- **Hallucination Rate:** Fraction of answers scoring ≤ 2 (lower is better).")
    lines.append("- **Execution Time:** Average latency per query in seconds (if available).")
    lines.append("")

    # 4. Pipeline Description
    lines.append("## 4. RAG Pipeline Description")
    lines.append("")
    lines.append(
        "Each model operates within the same RAG pipeline:"
    )
    lines.append("")
    lines.append("1. **Retrieval:** Query embedding and top-k (k=3) document retrieval from ChromaDB.")
    lines.append("2. **Context Augmentation:** Retrieved documents concatenated into a context string.")
    lines.append("3. **Generation:** LLM generates answer given the query and context.")
    lines.append("4. **Evaluation:** Judge LLM scores the answer against the reference.")
    lines.append("")
    lines.append("All models use the same embedding model (sentence-transformers:all-MiniLM-L6-v2) "
                 "to ensure fair comparison at the retrieval stage.")
    lines.append("")

    # 5. Results and Evaluation
    lines.append("## 5. Results and Evaluation")
    lines.append("")

    # Table of results
    lines.append("### 5.1 Model Comparison Table")
    lines.append("")
    lines.append("| Model | Variant | Avg Score | Hallucination Rate | Items Evaluated |")
    lines.append("|-------|---------|-----------|------------------|-----------------|")

    for r in sorted(results, key=lambda x: x["avg_score"], reverse=True):
        lines.append(
            f"| {r['model_name']} | {r.get('variant', 'default')} | "
            f"{r['avg_score']:.2f}/5 | {r['hallucination_rate']*100:.1f}% | {r['num_items']} |"
        )

    lines.append("")

    # Visualizations
    lines.append("### 5.2 Visual Analysis")
    lines.append("")
    lines.append("![Accuracy Comparison](graphs/01_accuracy_comparison.png)")
    lines.append("")
    lines.append("![Hallucination Comparison](graphs/02_hallucination_comparison.png)")
    lines.append("")
    lines.append("![Score Distribution](graphs/03_score_distribution.png)")
    lines.append("")
    lines.append("![Accuracy vs Hallucination Trade-off](graphs/04_accuracy_vs_hallucination.png)")
    lines.append("")

    # Score distribution breakdown
    lines.append("### 5.3 Detailed Score Distribution")
    lines.append("")
    for r in sorted(results, key=lambda x: x["avg_score"], reverse=True):
        lines.append(f"**{r['model_name']}**")
        dist = r.get("score_distribution", {})
        for score in range(5, 0, -1):
            count = dist.get(score, 0)
            pct = (count / r['num_items'] * 100) if r['num_items'] > 0 else 0
            lines.append(f"- Score {score}: {count} items ({pct:.1f}%)")
        lines.append("")

    lines.append("### 5.4 Analysis and Findings")
    lines.append("")

    # Find best, worst, and most reliable
    best = max(results, key=lambda x: x["avg_score"])
    worst = min(results, key=lambda x: x["avg_score"])
    most_reliable = min(results, key=lambda x: x["hallucination_rate"])

    lines.append(f"**Best Performing Model:** `{best['model_name']}` with avg score {best['avg_score']:.2f}/5")
    lines.append("")
    lines.append(
        f"**Most Reliable (lowest hallucination):** `{most_reliable['model_name']}` "
        f"with {most_reliable['hallucination_rate']*100:.1f}% hallucination rate"
    )
    lines.append("")
    lines.append(f"**Least Accurate:** `{worst['model_name']}` with avg score {worst['avg_score']:.2f}/5")
    lines.append("")

    # Comparative commentary
    lines.append("#### Key Observations:")
    lines.append("")

    accuracy_range = best["avg_score"] - worst["avg_score"]
    lines.append(f"- **Accuracy Range:** {accuracy_range:.2f} points ({worst['avg_score']:.2f} to {best['avg_score']:.2f})")

    hallucination_range = max(r["hallucination_rate"] for r in results) - min(r["hallucination_rate"] for r in results)
    lines.append(f"- **Hallucination Rate Range:** {hallucination_range*100:.1f}% across models")

    high_hallucination = [r for r in results if r["hallucination_rate"] > 0.3]
    if high_hallucination:
        models_str = ", ".join([f"`{r['model_name']}`" for r in high_hallucination])
        lines.append(f"- **High Hallucination Risk:** {models_str} show >30% hallucination rates")

    reliable = [r for r in results if r["hallucination_rate"] < 0.15]
    if reliable:
        models_str = ", ".join([f"`{r['model_name']}`" for r in reliable])
        lines.append(f"- **Recommended (low hallucination):** {models_str}")

    lines.append("")

    # 6. Conclusion
    lines.append("## 6. Conclusion")
    lines.append("")
    lines.append(
        f"Based on this evaluation, `{best['model_name']}` achieves the highest accuracy "
        f"with {best['avg_score']:.2f}/5. For production use prioritizing reliability, "
        f"`{most_reliable['model_name']}` is recommended due to its lowest hallucination rate "
        f"({most_reliable['hallucination_rate']*100:.1f}%). "
        f"The trade-off between accuracy and hallucination risk should guide model selection "
        f"based on application requirements."
    )
    lines.append("")

    return "\n".join(lines)


def create_docx_report(results: List[Dict], dataset_stats: Dict, md_content: str) -> Document:
    """Generate comprehensive DOCX report."""
    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    # Title
    title = doc.add_heading("RAG + LLM Evaluation Report", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_paragraph("Comprehensive Comparison of Language Models in Retrieval-Augmented Generation")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_format = subtitle.runs[0]
    subtitle_format.italic = True

    doc.add_paragraph(f"Date: {datetime.now().strftime('%B %d, %Y')}", style="Normal")
    doc.add_paragraph("")

    # 1. Introduction
    doc.add_heading("1. Introduction", level=1)
    doc.add_paragraph(
        "This report presents a comprehensive evaluation of multiple Large Language Models (LLMs) "
        "in a Retrieval-Augmented Generation (RAG) pipeline. The goal is to compare model performance "
        "in terms of accuracy (via LLM-as-a-judge scoring), hallucination rates, and execution time. "
        "This analysis helps identify which models provide the best balance between accuracy and "
        "reliability for production RAG deployments."
    )

    # 2. Dataset Description
    doc.add_heading("2. Dataset Description", level=1)
    doc.add_paragraph(
        f"The evaluation uses a curated QA dataset derived from financial and banking domain documents."
    )
    
    dataset_table = doc.add_table(rows=3, cols=2)
    dataset_table.style = "Light Grid Accent 1"
    dataset_table.cell(0, 0).text = "Metric"
    dataset_table.cell(0, 1).text = "Value"
    dataset_table.cell(1, 0).text = "Total QA Pairs"
    dataset_table.cell(1, 1).text = str(dataset_stats.get('total_qa_pairs', 'N/A'))
    dataset_table.cell(2, 0).text = "Unique Documents"
    dataset_table.cell(2, 1).text = str(dataset_stats.get('unique_docs', 'N/A'))

    doc.add_paragraph("Domain: Banking, Insurance, Loans, and Financial Services")
    doc.add_paragraph("Question Types: Factual recall, comparison, definition, numerical extraction")
    doc.add_paragraph(
        "Each QA pair includes a reference answer manually curated from source documents, "
        "ensuring ground-truth labels for evaluation."
    )

    # 3. Method Description
    doc.add_heading("3. Method Description", level=1)
    doc.add_heading("3.1 Evaluation Methodology", level=2)
    doc.add_paragraph(
        "We employ an LLM-as-a-Judge evaluation approach, where a secondary LLM scores each model's "
        "answer against the reference answer on a scale of 1–5:"
    )

    scoring_table = doc.add_table(rows=6, cols=2)
    scoring_table.style = "Light Grid Accent 1"
    scoring_table.cell(0, 0).text = "Score"
    scoring_table.cell(0, 1).text = "Description"
    scoring_table.cell(1, 0).text = "5"
    scoring_table.cell(1, 1).text = "Fully correct and complete"
    scoring_table.cell(2, 0).text = "4"
    scoring_table.cell(2, 1).text = "Mostly correct, minor omissions or imprecision"
    scoring_table.cell(3, 0).text = "3"
    scoring_table.cell(3, 1).text = "Partially correct, key ideas present but incomplete"
    scoring_table.cell(4, 0).text = "2"
    scoring_table.cell(4, 1).text = "Mostly incorrect, minimal relevant content"
    scoring_table.cell(5, 0).text = "1"
    scoring_table.cell(5, 1).text = "Incorrect or irrelevant"

    doc.add_paragraph(
        "Answers with scores ≤ 2 are classified as hallucinations or serious errors, "
        "as they deviate significantly from the reference."
    )

    doc.add_heading("3.2 Key Metrics", level=2)
    doc.add_paragraph("Average Score (1–5): Mean LLM-as-a-judge score across all QA pairs.", style="List Bullet")
    doc.add_paragraph("Hallucination Rate: Fraction of answers scoring ≤ 2 (lower is better).", style="List Bullet")
    doc.add_paragraph("Execution Time: Average latency per query in seconds (if available).", style="List Bullet")

    # 4. Pipeline Description
    doc.add_heading("4. RAG Pipeline Description", level=1)
    doc.add_paragraph("Each model operates within the same RAG pipeline:")
    doc.add_paragraph("Query embedding and top-k (k=3) document retrieval from ChromaDB.", style="List Number")
    doc.add_paragraph("Retrieved documents concatenated into a context string.", style="List Number")
    doc.add_paragraph("LLM generates answer given the query and context.", style="List Number")
    doc.add_paragraph("Judge LLM scores the answer against the reference.", style="List Number")

    doc.add_paragraph(
        "All models use the same embedding model (sentence-transformers:all-MiniLM-L6-v2) "
        "to ensure fair comparison at the retrieval stage."
    )

    # 5. Results and Evaluation
    doc.add_heading("5. Results and Evaluation", level=1)

    doc.add_heading("5.1 Model Comparison Table", level=2)
    results_table = doc.add_table(rows=1, cols=5)
    results_table.style = "Light Grid Accent 1"
    hdr = results_table.rows[0].cells
    hdr[0].text = "Model"
    hdr[1].text = "Variant"
    hdr[2].text = "Avg Score"
    hdr[3].text = "Hallucination Rate"
    hdr[4].text = "Items"

    for r in sorted(results, key=lambda x: x["avg_score"], reverse=True):
        row = results_table.add_row().cells
        row[0].text = str(r["model_name"])
        row[1].text = str(r.get("variant", "default"))
        row[2].text = f"{r['avg_score']:.2f}/5"
        row[3].text = f"{r['hallucination_rate']*100:.1f}%"
        row[4].text = str(r["num_items"])

    doc.add_heading("5.2 Visual Analysis", level=2)
    doc.add_paragraph("Below are visualizations comparing model performance across key metrics:")
    
    if (GRAPHS_DIR / "01_accuracy_comparison.png").exists():
        doc.add_paragraph("Accuracy Comparison:")
        doc.add_picture(str(GRAPHS_DIR / "01_accuracy_comparison.png"), width=Inches(6))
    
    if (GRAPHS_DIR / "02_hallucination_comparison.png").exists():
        doc.add_paragraph("Hallucination Rate Comparison:")
        doc.add_picture(str(GRAPHS_DIR / "02_hallucination_comparison.png"), width=Inches(6))
    
    if (GRAPHS_DIR / "03_score_distribution.png").exists():
        doc.add_paragraph("Score Distribution:")
        doc.add_picture(str(GRAPHS_DIR / "03_score_distribution.png"), width=Inches(6))
    
    if (GRAPHS_DIR / "04_accuracy_vs_hallucination.png").exists():
        doc.add_paragraph("Accuracy vs. Hallucination Trade-off:")
        doc.add_picture(str(GRAPHS_DIR / "04_accuracy_vs_hallucination.png"), width=Inches(6))

    doc.add_heading("5.3 Detailed Score Distribution", level=2)
    for r in sorted(results, key=lambda x: x["avg_score"], reverse=True):
        doc.add_paragraph(f"{r['model_name']}", style="Heading 3")
        dist = r.get("score_distribution", {})
        for score in range(5, 0, -1):
            count = dist.get(score, 0)
            pct = (count / r['num_items'] * 100) if r['num_items'] > 0 else 0
            doc.add_paragraph(f"Score {score}: {count} items ({pct:.1f}%)", style="List Bullet")

    doc.add_heading("5.4 Analysis and Findings", level=2)
    best = max(results, key=lambda x: x["avg_score"])
    worst = min(results, key=lambda x: x["avg_score"])
    most_reliable = min(results, key=lambda x: x["hallucination_rate"])

    doc.add_paragraph(f"Best Performing Model: {best['model_name']} with avg score {best['avg_score']:.2f}/5")
    doc.add_paragraph(
        f"Most Reliable (lowest hallucination): {most_reliable['model_name']} "
        f"with {most_reliable['hallucination_rate']*100:.1f}% hallucination rate"
    )
    doc.add_paragraph(f"Least Accurate: {worst['model_name']} with avg score {worst['avg_score']:.2f}/5")

    doc.add_heading("Key Observations:", level=3)
    accuracy_range = best["avg_score"] - worst["avg_score"]
    doc.add_paragraph(
        f"Accuracy Range: {accuracy_range:.2f} points ({worst['avg_score']:.2f} to {best['avg_score']:.2f})",
        style="List Bullet"
    )

    hallucination_range = max(r["hallucination_rate"] for r in results) - min(r["hallucination_rate"] for r in results)
    doc.add_paragraph(
        f"Hallucination Rate Range: {hallucination_range*100:.1f}% across models",
        style="List Bullet"
    )

    high_hallucination = [r for r in results if r["hallucination_rate"] > 0.3]
    if high_hallucination:
        models_str = ", ".join([f"{r['model_name']}" for r in high_hallucination])
        doc.add_paragraph(f"High Hallucination Risk: {models_str} show >30% hallucination rates", style="List Bullet")

    reliable = [r for r in results if r["hallucination_rate"] < 0.15]
    if reliable:
        models_str = ", ".join([f"{r['model_name']}" for r in reliable])
        doc.add_paragraph(f"Recommended (low hallucination): {models_str}", style="List Bullet")

    # 6. Conclusion
    doc.add_page_break()
    doc.add_heading("6. Conclusion", level=1)
    doc.add_paragraph(
        f"Based on this evaluation, {best['model_name']} achieves the highest accuracy "
        f"with {best['avg_score']:.2f}/5. For production use prioritizing reliability, "
        f"{most_reliable['model_name']} is recommended due to its lowest hallucination rate "
        f"({most_reliable['hallucination_rate']*100:.1f}%). "
        f"The trade-off between accuracy and hallucination risk should guide model selection "
        f"based on application requirements."
    )

    return doc


def main() -> None:
    # Load evaluation results
    eval_files: Dict[str, Dict[str, Path]] = {
        "gemma3:1b": {
            "default": ARTIFACTS_DIR / "gemma3_eval.jsonl",
        },
        "qwen:latest": {
            "default": ARTIFACTS_DIR / "qwen_eval.jsonl",
        },
        "Phi3:latest": {
            "default": ARTIFACTS_DIR / "phi3_eval.jsonl",
        },
        "deepseek-r1:1.5b": {
            "default": ARTIFACTS_DIR / "deepseek_eval.jsonl",
        },
        "smollm:latest": {
            "default": ARTIFACTS_DIR / "smolLM_eval.jsonl",
        },
    }

    results: List[Dict] = []
    for model_name, variants in eval_files.items():
        for variant, path in variants.items():
            if not path.exists():
                continue
            summary = summarize_eval(path)
            summary.update({
                "model_name": model_name,
                "variant": variant,
            })
            results.append(summary)

    if not results:
        print("No evaluation files found!")
        return

    # Get dataset stats
    dataset_stats = get_dataset_stats()

    # Generate graphs
    print("Generating graphs...")
    graph_paths = generate_graphs(results)
    print(f"✓ {len(graph_paths)} graphs generated in {GRAPHS_DIR}")

    # Generate markdown report
    print("Generating markdown report...")
    md_content = create_markdown_report(results, dataset_stats)
    REPORT_MD_PATH.parent.mkdir(parents=True, exist_ok=True)
    REPORT_MD_PATH.write_text(md_content, encoding="utf-8")
    print(f"✓ Markdown report written to {REPORT_MD_PATH}")

    # Generate DOCX report
    print("Generating DOCX report...")
    doc = create_docx_report(results, dataset_stats, md_content)
    REPORT_DOCX_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(REPORT_DOCX_PATH)
    print(f"✓ DOCX report written to {REPORT_DOCX_PATH}")

    print(f"\n✓ Report Summary:")
    print(f"  - Total models evaluated: {len(results)}")
    print(f"  - Total QA pairs: {dataset_stats['total_qa_pairs']}")
    print(f"  - Unique documents: {dataset_stats['unique_docs']}")
    print(f"  - Output location: {ARTIFACTS_DIR}")


if __name__ == "__main__":
    main()
